#!/usr/bin/env python3
"""
Bouwt het NESA V3 draaiboek als .docx.

Verwerkt alle feedback:
  - 9 marge-opmerkingen Evi (V2.1)
  - Infographic 'Wat is NESA?' (V2.1)
  - Gespreknotities Sam ↔ Evi 26-05 (V3): concrete cijfers, behandelaars,
    tone-of-voice meer informerend dan marketing, intake-flow uitgewerkt.

Geen TODO-flags meer — alles is concreet ingevuld.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/Users/samvanmourik/Downloads/Draaiboek NESA V3.docx"

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def H1(text): return doc.add_heading(text, level=1)
def H2(text): return doc.add_heading(text, level=2)
def H3(text): return doc.add_heading(text, level=3)

def para(text="", bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        if text:
            p.add_run(" " + text)
    elif text:
        p.add_run(text)
    return p

def voice_over(text):
    p = doc.add_paragraph()
    run = p.add_run("Voice-over: ")
    run.bold = True
    italic = p.add_run('"' + text + '"')
    italic.italic = True
    return p

def bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if "**" in item:
            parts = item.split("**")
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True
        else:
            p.add_run(item)

def labeled_block(label, text):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    if text:
        p.add_run("\n" + text)

def scene_header(num, time, title):
    H2(f"Scene {num} · {time} · {title}")

# ═══════════════════════════════════════════════════════════
# DOCUMENT
# ═══════════════════════════════════════════════════════════

doc.add_heading("Draaiboek animatie explainer video — V3", level=0)

# ── Projectoverzicht ──
H1("Projectoverzicht")
tbl = doc.add_table(rows=4, cols=2)
tbl.style = "Light Grid Accent 1"
rows = [
    ("Opdrachtgever", "NESAclinics"),
    ("Format", "Animatie explainer video"),
    ("Duur", "Circa 1:10 (richtbereik 1:00 — 1:30)"),
    ("Aantal scenes", "10"),
]
for i, (l, r) in enumerate(rows):
    tbl.rows[i].cells[0].text = l
    tbl.rows[i].cells[1].text = r

para()
labeled_block("Doelgroep",
    "Mensen (30+) met langdurige klachten die al een regulier zorgtraject (huisarts, fysio, ziekenhuis) hebben doorlopen zonder voldoende resultaat. Op het moment dat ze deze video zien zijn ze al geïnformeerd over NESA via brochure of website — de video wordt getoond op informatie-momenten of op de site.")

labeled_block("Doel video",
    "Informeren (geen cold marketing). De kijker is al benieuwd; deze video moet vertrouwen geven door concreet uit te leggen wat NESA is, hoe het werkt, wie het uitvoert, wat het kost en hoe je begint. De CTA is het plannen van een intake.")

labeled_block("Animatiestijl",
    "Cartoon-flat illustraties: dikke donkere outlines, vlakke kleurvlakken, ronde gezichten met zachte blosjes, subtiele schaduwen onder figuren. Karakter-gedreven, weinig afleidende elementen op de achtergrond — clean.")

labeled_block("Kleurpalet",
    "NESAclinics huisstijl. Achtergrond door alle scenes neutraal beige #ECEDE4. Karakters en accenten in mosgroen #687959, oudroze #D67489 en lichtere varianten. Eindframe wit voor merkpresentatie.")

labeled_block("Tone of voice",
    "Rustig, geloofwaardig, helder. Erkent teleurstelling van eerdere trajecten zonder reguliere zorg af te kraken. Geen overdreven beloften, geen pijn-druk, geen verkopen — wel duidelijke uitleg. Wat de kijker kan verwachten staat centraal.")

# ── Persona's ──
H1("Drie persona's (op basis van NESAclinics aanmeldingsdata)")
para("Komen terug in scene 1 (klachten) en scene 8 (hersteld). De drie persona's vertegenwoordigen de drie grootste klachtgroepen.")

H3("Peter, 63 — chronische pijnklachten / neuropathie (PH)")
labeled_block("Visueel:", "Man (60+), grijze baard met bril, eenvoudige spencer.")
labeled_block("Klachten in beeld (scene 1):", "Pijn-streepjes bij onderrug, hand op heup, gespannen schouders.")

H3("Marieke, 47 — chronische vermoeidheid & slapeloosheid (MV)")
labeled_block("Visueel:", "Vrouw (40-50), donker haar opgestoken, eenvoudige bloes.")
labeled_block("Klachten in beeld (scene 1):", "Donkere kringen onder de ogen, lichte vermoeidheidsstreepjes rond hoofd, hand tegen voorhoofd.")

H3("Sophie, 34 — long covid / brainfog (SP)")
labeled_block("Visueel:", "Vrouw (30-40), bruin haar in dutje, casual top.")
labeled_block("Klachten in beeld (scene 1):", "Hand tegen voorhoofd, kleine 'mistwolkjes' rond hoofd (brainfog), zit gebogen op bankrand.")

# ═══════════════════════════════════════════════════════════
# SCENES
# ═══════════════════════════════════════════════════════════

H1("Scene-voor-scene draaiboek")

# ──────── SCENE 1 ────────
scene_header(1, "0:00 — 0:08", "OPENING — Herkenning")
voice_over("Pijn die niet weggaat. Vermoeidheid die nooit verdwijnt. Een lichaam dat niet meer tot rust komt. Herkenbaar?")
labeled_block("Beeldbeschrijving",
    "Drie karakters naast elkaar, op enige afstand. Cartoon-flat stijl, dikke outlines, ronde gezichten.")
bullets([
    "**Links: Peter** (63, grijze baard, bril, spencer) — gefronste blik, hand op onderrug, kleine pijn-streepjes bij rug en bovenbeen.",
    "**Midden: Marieke** (47, opgestoken donker haar, bloes) — donkere kringen onder de ogen, hand tegen voorhoofd, slappe schouders, kleine 'zzz'- of vermoeidheidsstreepjes rond hoofd.",
    "**Rechts: Sophie** (34, bruin haar in dutje, casual top) — hand tegen voorhoofd, kleine mistwolkjes rond hoofd (brainfog), zit gebogen op bankrand.",
])
para("Onder elke figuur een subtiele ovale schaduw.")
labeled_block("Animatie:", "karakters fade'n één voor één in (0.4 sec interval). Lichte idle-animatie per karakter (rustige ademhaling, een hand die kort beweegt, mistwolkjes die langzaam pulseren).")
labeled_block("Kleurgebruik:", "")
bullets([
    "Achtergrond: neutraal beige #ECEDE4",
    "Outlines: donker #353D38 (3-4px equivalent)",
    "Peter: huidskleur natuurlijk, spencer mosgroen licht #A8B195, broek beige, baard grijs",
    "Marieke: huidskleur natuurlijk, bloes oudroze licht #F196AC, rok donker",
    "Sophie: huidskleur natuurlijk, top wit #FFFFFF, broek grijs-bruin",
    "Pijn-streepjes: donker #353D38",
    "Vermoeidheidsstreepjes en mistwolkjes: mosgroen licht #A8B195 met lage opacity",
])
labeled_block("Typografie:", "geen tekst on-screen.")

# ──────── SCENE 2 ────────
scene_header(2, "0:08 — 0:14", "PROBLEEM — Wat al geprobeerd is")
voice_over("Je hebt al veel geprobeerd. Huisarts, fysio, ziekenhuis, onderzoek na onderzoek. Maar de klachten blijven.")
labeled_block("Beeldbeschrijving",
    "Eén centraal karakter — Marieke werkt hier goed visueel. Rondom haar verschijnen drie iconen op witte cirkels: stethoscoop (huisarts), handpalm/massage (fysio), kruis (ziekenhuis). Elk krijgt na 0.5 sec een grijs vinkje ('gehad').")
para("Tot slot fade'n de iconen weg en blijft alleen een denkwolk met een grijs vraagteken boven Marieke's hoofd, die zacht pulseert.")
labeled_block("Animatie:", "iconen verschijnen één voor één in cirkel rondom (0.5 sec elk), vinkjes draaien in. Op het einde verdwijnen de iconen en blijft alleen het vraagteken.")
labeled_block("Kleurgebruik:", "")
bullets([
    "Achtergrond: neutraal beige #ECEDE4",
    "Outlines: donker #353D38",
    "Karakter-kleuren: consistent met scene 1",
    "Icoon-cirkels: wit #FFFFFF met donkere outline",
    "Iconen zelf: mosgroen #687959",
    "Vinkjes: grijs #7E7E7E",
    "Vraagteken-wolk: wit #FFFFFF met donkere outline",
])
labeled_block("Typografie:", "geen tekst on-screen.")

# ──────── SCENE 3 ────────
scene_header(3, "0:14 — 0:21", "INZICHT — Centraal zenuwstelsel uit balans")
voice_over("Bij langdurige klachten raakt je centrale zenuwstelsel uit balans. Je lichaam blijft in 'stand-by' staan — herstel komt niet meer op gang.")
labeled_block("Beeldbeschrijving",
    "Centraal: een eenvoudig lichaam-silhouet (van voren, halftransparant, geen detailgezicht).")
para("In het silhouet duidelijk zichtbaar: hersenen bovenaan, ruggenmerg dat naar beneden loopt (centraal zenuwstelsel), met fijne vertakkingen naar armen, benen en romp (perifere zenuwen).")
para("Rond de hersenen een lichte gloed-overlay in oudroze, die onrustig pulseert — visualisatie van 'overactief / in stand-by'.")
para("Onder het silhouet een eenvoudige balans-balk met links 'Stand-by' en rechts 'Herstel'. De balans helt zichtbaar over naar links.")
labeled_block("Animatie:", "silhouet fade'n in. Zenuwstelsel tekent zichzelf vanaf hersenen naar buiten (centraal → perifeer). Gloed rond hersenen pulseert snel en onrustig. Balans helt naar links in 2 sec.")
labeled_block("Kleurgebruik:", "")
bullets([
    "Achtergrond: neutraal beige #ECEDE4",
    "Lichaam-silhouet: mosgroen licht #A8B195 met lage opacity, donkere outline #353D38",
    "Centraal zenuwstelsel (hersenen + ruggenmerg): donker #353D38 met fijne lijnvoering",
    "Perifere zenuwen (vertakkingen): oudroze #D67489, dun",
    "Hersenen-gloed 'overactief': oudroze #D67489 met zachte glow, onrustig pulserend",
    "Balans-balk en wijzer: donker #353D38",
    "Labels: donker #353D38",
])
labeled_block("Typografie:", "'Stand-by' en 'Herstel' in Open Sans Bold, donker #353D38.")

# ──────── SCENE 4 ────────
scene_header(4, "0:21 — 0:28", "OPLOSSING — Wat NESA-therapie is")
voice_over("NESA-therapie helpt je systeem weer in balans te komen. Met zeer zachte, pijnloze prikkels via je perifere zenuwen. Niet-invasief — geen medicijnen, geen naalden.")
labeled_block("Beeldbeschrijving",
    "Apparaat-weergave conform de werkelijke NESA-foto: een wit, rechthoekig kastje met een bedieningspaneel en vijf gekleurde aansluit-poorten onderaan (zwart, geel, groen, blauw, rood).")
bullets([
    "Centraal: het **NESA-kastje** (wit, rechthoekig). Bovenkant met klein donker display-vlakje, daaronder een paneel met enkele schakelaars. Onderaan vijf ronde aansluitingen in herkenbare kleuren.",
    "Vanaf de gekleurde aansluitingen lopen dunne kabels naar **handschoenen** (polsband-stijl, zonder vingers) en **sokken** (lage kous) die naast het kastje liggen.",
    "Rechts in beeld drie kleine icoon-callouts die één voor één in fade'n synchroon met de voice-over:",
])
bullets([
    "Blaadje-icoon — *'Pijnloos'*",
    "Schild-icoon — *'Niet-invasief'*",
    "Hersenen-icoon — *'Ondersteunt zelfherstel'*",
])
labeled_block("Animatie:", "kastje fade'n in centraal. Vanuit elke gekleurde aansluiting reist een rustige pulse langs de kabel naar buiten (richting handschoen/sok). Icoon-callouts verschijnen rechts één voor één (0.6 sec interval). Géén vonkjes, géén bliksem-symbolen.")
labeled_block("Kleurgebruik:", "")
bullets([
    "Achtergrond: neutraal beige #ECEDE4",
    "Kastje: wit #FFFFFF met donkere #353D38 outline en lichtgrijze schaduwvlakken",
    "Display-vlakje op kastje: donkergrijs (off-state) of mosgroen licht #A8B195 (on-state)",
    "Aansluit-poorten onderaan: NESA-kleuren — zwart, geel #F0C04F, groen #687959, blauw #5B7D9A, rood #C25A5A",
    "Kabels: donker #353D38, dun",
    "Handschoenen en sokken: wit #FFFFFF met donkere outline en mosgroen #687959 band",
    "Elektrode-markers op handschoenen/sokken: kleine oudroze #D67489 cirkels",
    "Pulse-animatie langs kabels: oudroze licht #F196AC met zachte glow",
    "Icoon-callouts: wit #FFFFFF cirkel met donkere outline, icoon in mosgroen #687959, labeltekst donker #353D38",
])
labeled_block("Typografie:", "icoon-labels in Open Sans Bold klein, donker #353D38.")

# ──────── SCENE 5 ────────
scene_header(5, "0:28 — 0:36", "WERKING — Van prikkel naar herstel")
voice_over("De milde prikkels reizen via je zenuwbanen naar je hersenstam, en verspreiden zich naar belangrijke hersengebieden. Daar zetten ze een natuurlijk herstelproces in gang — je lichaam zoekt actief naar nieuwe balans.")
labeled_block("Beeldbeschrijving",
    "Twee-luik compositie:")
bullets([
    "**Links:** half-liggend lichaam-silhouet (zelfde stijl als scene 3, in profiel of driekwart). Vanaf polsen en enkels vloeit een rustige pulse-stroom door de zenuwbanen omhoog, samenkomend bij de wervelkolom en eindigend bij de hersenstam.",
    "**Rechts:** illustratieve weergave van de hersenen (zijaanzicht, vereenvoudigd). Vier gebieden lichten één voor één op, elk in eigen kleur, met een klein labeltje:",
])
bullets([
    "**Cortex** — paars-grijs (waarneming en gevoel)",
    "**Limbisch systeem** — oudroze (emoties, stressregulatie)",
    "**Cerebellum** — mosgroen (motoriek, balans)",
    "**Hersenstam** — beige (ademhaling, hartslag, slaap)",
])
para("Onderaan een kleine cyclische visualisatie: drie cirkels in een lus met pijlen, met labels 'Signalen → Hormonen → Lichaam reageert'. Toont het zelfregulerende systeem.")
para("Daaronder de balans-balk uit scene 3, met de wijzer die nu langzaam van 'Stand-by' naar 'Herstel' schuift.")
labeled_block("Animatie:", "pulse-stroom start bij polsen en enkels (continue beweging omhoog). Hersengebieden lichten één voor één op (cortex → limbisch → cerebellum → hersenstam, 0.6 sec interval) synchroon met voice-over. Lus-visualisatie rolt eenmaal door. Balans-wijzer schuift in 3 sec naar rechts.")
labeled_block("Kleurgebruik:", "")
bullets([
    "Achtergrond: neutraal beige #ECEDE4 met subtiele lichtere gloed rond hersenen",
    "Outlines: donker #353D38",
    "Lichaam-silhouet: mosgroen licht #A8B195 vulling, donkere outline",
    "Pulse-stroom door zenuwbanen: oudroze #D67489 lijn met oudroze licht #F196AC glow",
    "Hersenen-omtrek: donker #353D38",
    "Cortex-highlight: paars-grijs #8E89A8",
    "Limbisch systeem-highlight: oudroze #D67489",
    "Cerebellum-highlight: mosgroen #687959",
    "Hersenstam-highlight: beige #C9B179",
    "Labels hersengebieden: donker #353D38 in Open Sans Bold klein",
    "Lus-cirkels: mosgroen licht #A8B195 met donkere outline, pijlen donker #353D38",
    "Labels 'Signalen', 'Hormonen', 'Lichaam reageert': donker #353D38",
    "Balans-balk: donker #353D38, wijzer schuift naar rechts",
])
labeled_block("Typografie:", "alle labels in Open Sans Bold klein, donker #353D38. 'Stand-by' en 'Herstel' op balans-balk in Open Sans Regular.")

# ──────── SCENE 6 ────────
scene_header(6, "0:36 — 0:43", "SESSIE — Hoe een behandeling eruitziet")
voice_over("Een sessie duurt ongeveer een uur. Je ligt ontspannen op een behandelbank, met kleine elektroden op je polsen, enkels en nek. De signalen zijn nauwelijks voelbaar en volledig pijnvrij.")
labeled_block("Beeldbeschrijving", "Brede shot:")
bullets([
    "Links in beeld: het NESA-kastje (zelfde als scene 4) op een verrijdbaar tafeltje. Kabels lopen naar de patiënt.",
    "Midden: een **liggende patiënt** (Marieke, voor continuïteit). Ze ligt op een halfschuine behandelbank — niet plat. Kussen onder hoofd, deken half over benen. Ogen ontspannen dicht, rustige glimlach. Handen rusten op buik in handschoenen, voeten in sokken, kleine elektrode-patches zichtbaar op polsen, enkels en in de nek (kleine cirkels met dunne lijn).",
    "Rechts: de therapeut (rond de 40, mosgroen polo) staat naast haar en kijkt rustig op het kastje.",
])
para("Subtiele pulse-animaties bij polsen, enkels en op het kastje. Géén vonkjes, géén elektriciteit-symbolen.")
labeled_block("Animatie:", "scene fade'n in. Patiënt-lichaam ademt rustig (subtiele heffing buik). Pulse-animaties starten zodra therapeut achter het kastje plaatsneemt.")
labeled_block("Kleurgebruik:", "")
bullets([
    "Achtergrond: neutraal beige #ECEDE4",
    "Outlines: donker #353D38",
    "Kastje: wit #FFFFFF zoals scene 4",
    "Behandelbank: licht grijs-bruin met donkere outline",
    "Kussen: wit #FFFFFF",
    "Deken: mosgroen licht #A8B195",
    "Patiënt: kleuren consistent met scene 1",
    "Handschoenen en sokken: wit #FFFFFF met mosgroen #687959 band",
    "Elektrode-patches: huidkleurig met donkere outline, kleine oudroze #D67489 pulse-cirkel eromheen",
    "Therapeut: huidskleur, polo mosgroen #687959, broek grijs-bruin",
])
labeled_block("Typografie:", "geen tekst on-screen.")

# ──────── SCENE 7 ────────
scene_header(7, "0:43 — 0:49", "BEHANDELAARS — Wie de behandeling uitvoert")
voice_over("De behandeling wordt gegeven door medisch geschoolde professionals — fysiotherapeuten, osteopaten of artsen, allemaal speciaal opgeleid in NESA-therapie. Op meer dan 40 locaties in Nederland.")
labeled_block("Beeldbeschrijving",
    "Drie behandelaars naast elkaar in beeld, schouder aan schouder maar elk met een eigen kleine context-prop. Cartoon-flat stijl, dezelfde lijnvoering als alle andere scenes.")
bullets([
    "**Links: fysiotherapeut** — vrouw, ongeveer 40, mosgroen polo. Houdt subtiel een handschoen-element vast of heeft hand op de schouder van een denkbeeldige patiënt. Klein label onder figuur: 'Fysiotherapeut'.",
    "**Midden: osteopaat** — man, ongeveer 45, mosgroen polo, iets formeler. Klein label: 'Osteopaat'.",
    "**Rechts: arts** — vrouw, ongeveer 50, wit-met-mosgroen accenten (subtiele verwijzing naar witte jas zonder letterlijke witte jas). Klein label: 'Arts'.",
])
para("Onder de drie figuren een dunne tijdlijn of horizontale lijn met centraal een NESA-badge (een klein rond logo of icoon + tekst 'Opgeleid in NESA-therapie'). De badge verbindt alle drie de figuren visueel.")
para("Rechts naast (of onder) de drie figuren een kleine illustratie: contour van Nederland met meerdere stippen erop verdeeld, met daaronder '40+ locaties'.")
labeled_block("Animatie:", "drie behandelaars fade'n één voor één in (0.4 sec interval). NESA-badge verschijnt centraal met pop-in. Nederland-kaart fade'n in rechts; stippen verschijnen in golf van zuid naar noord (~0.05 sec interval).")
labeled_block("Kleurgebruik:", "")
bullets([
    "Achtergrond: neutraal beige #ECEDE4",
    "Outlines: donker #353D38",
    "Behandelaars: huidskleuren natuurlijk, mosgroen #687959 polo's, broeken grijs-bruin",
    "Label-balkjes onder figuren: wit #FFFFFF met donkere #353D38 tekst",
    "NESA-badge: wit #FFFFFF cirkel met mosgroen #687959 rand, tekst donker #353D38",
    "Nederland-kaart: wit #FFFFFF vulling met donkere outline",
    "Locatie-stippen: oudroze #D67489 met donkere outline",
    "'40+ locaties': mosgroen #687959 bold, met dunne 'in Nederland' eronder in donker #353D38",
])
labeled_block("Typografie:", "labels onder figuren in Open Sans Regular, donker #353D38. NESA-badge tekst en '40+ locaties' in Open Sans Bold.")

# ──────── SCENE 8 ────────
scene_header(8, "0:49 — 0:56", "RESULTAAT — Wat veel mensen ervaren")
voice_over("De meeste mensen merken al binnen drie tot vijf sessies de eerste verbetering. Rustiger slapen, minder pijn, meer energie en focus.")
labeled_block("Beeldbeschrijving",
    "Drie karakters naast elkaar — Peter, Marieke, Sophie — nu in herstelde houding:")
bullets([
    "**Links: Peter** — staat rechtop, ontspannen schouders, geen hand meer op de rug, kleine glimlach. Pijn-streepjes weg.",
    "**Midden: Marieke** — donkere kringen weg, frisse blik, rechte houding, zelfverzekerd, kleine glimlach.",
    "**Rechts: Sophie** — mistwolkjes weg, rechtop staand met energieke houding.",
])
para("Boven elke figuur een klein lijn-icoon op witte cirkel:")
bullets([
    "Boven Peter: pijn-vrij icoon (een doorgekruiste pijnbliksem of een blije pijnschaal)",
    "Boven Marieke: maan met sterretjes (beter slapen)",
    "Boven Sophie: energie-symbool (volle batterij of zonnetje)",
])
para("Onder hen een dunne tijdlijn met drie markers: 'Sessie 1', 'Sessie 3', 'Sessie 5'.")
labeled_block("Animatie:", "karakters morphen van hun pijn-versie (scene 1) naar herstel-versie met zachte cross-fade van 0.6 sec. Iconen verschijnen één voor één met pop-in. Tijdlijn tekent zichzelf van links naar rechts.")
labeled_block("Kleurgebruik:", "")
bullets([
    "Achtergrond: neutraal beige #ECEDE4",
    "Karakters: identieke kleuren als scene 1 (continuïteit van herkenning)",
    "Iconen: oudroze #D67489 lijnwerk op witte #FFFFFF cirkel met donkere outline",
    "Tijdlijn: oudroze #D67489",
    "Sessie-labels: donker #353D38",
])
labeled_block("Typografie:", "sessie-labels in Open Sans Bold, donker #353D38.")

# ──────── SCENE 9 ────────
scene_header(9, "0:56 — 1:04", "HET TRAJECT — Hoe het na je aanmelding gaat")
voice_over("Na je aanmelding word je binnen zeven dagen gebeld om je intake te plannen. De intake kost 157 euro. Een volledig behandeltraject bestaat uit tien sessies en kost 925 euro. Veel mensen krijgen (een deel van) deze kosten vergoed.")
labeled_block("Beeldbeschrijving",
    "Een horizontale tijdlijn met vier stappen, elk een eenvoudig icoon op een witte cirkel met een korte label eronder:")
bullets([
    "**Stap 1 — Aanmelding:** envelop- of webformulier-icoon. Label: 'Aanmelding'.",
    "**Stap 2 — Belcontact:** telefoon-icoon met klein klokje of '7 dagen'-tekst. Label: 'Binnen 7 dagen contact'.",
    "**Stap 3 — Intake:** notitieblok-icoon. Label: 'Intake — € 157'.",
    "**Stap 4 — Traject:** een rij van 10 kleine bolletjes (sessies). Label: '10 sessies — € 925'.",
])
para("Onder de tijdlijn een klein dunner balkje met de tekst: 'Veel mensen krijgen (een deel van) de kosten vergoed — lees meer op nesaclinics.nl/kosten'.")
labeled_block("Animatie:", "tijdlijn tekent zichzelf van links naar rechts. Iconen poppen één voor één in (0.5 sec interval) synchroon met voice-over. De 10 sessie-bolletjes verschijnen één voor één (~0.1 sec elk) als kleine cascade.")
labeled_block("Kleurgebruik:", "")
bullets([
    "Achtergrond: neutraal beige #ECEDE4",
    "Tijdlijn: dunne donkere #353D38 lijn",
    "Stap-cirkels: wit #FFFFFF met donkere outline",
    "Iconen in cirkels: mosgroen #687959",
    "Bedrag-labels: donker #353D38 in Open Sans Bold",
    "Stap-labels: donker #353D38 in Open Sans Regular",
    "Sessie-bolletjes (stap 4): oudroze #D67489 vol",
    "Vergoeding-balkje onderaan: zacht mosgroen licht #A8B195 vulling, donkere tekst",
])
labeled_block("Typografie:", "bedrag-labels en hoofd-labels in Open Sans Bold, donker #353D38. Vergoeding-tekst onderaan in Open Sans Regular.")

# ──────── SCENE 10 ────────
scene_header(10, "1:04 — 1:10", "CTA — Plan je intake")
voice_over("Klaar om de eerste stap te zetten? Plan je intake op nesaclinics.nl.")
labeled_block("Beeldbeschrijving",
    "Eindframe (cross-fade van beige naar wit voor merkpresentatie):")
para("Centraal: NESAclinics logo in originele huisstijl-kleuren.")
para("Onder logo: tagline 'rust in je hoofd, kracht in je lijf'.")
para("Daaronder: CTA-knop in mosgroen #687959 met witte tekst 'Plan je intake'.")
para("Onder de knop: URL 'nesaclinics.nl' en rechts daarvan een QR-code.")
para("Subtiele huisstijl-lijnelementen (oudroze cirkels met streepjes, mosgroen lijnen) animeren in vanaf de zijkanten.")
labeled_block("Animatie:", "cross-fade beige → wit (0.5 sec). Elementen verschijnen in volgorde: logo → tagline → knop → URL/QR. Minimaal 2 seconden stilstaand zichtbaar.")
labeled_block("Kleurgebruik:", "")
bullets([
    "Achtergrond: wit #FFFFFF (eindframe-uitzondering)",
    "Logo: NESAclinics originele kleuren",
    "Tagline: donker #353D38",
    "CTA-knop: mosgroen #687959 vol, witte #FFFFFF tekst",
    "URL: donker #353D38",
    "QR-code: zwart op wit",
    "Huisstijl-lijnelementen: oudroze #D67489 cirkels met mosgroen #687959 lijnen",
])
labeled_block("Typografie:", "tagline in huisstijl-lettertype. CTA-knop in Open Sans Bold wit. URL in Open Sans Regular donker.")

# ── Productie ──
H1("Productienotities")
labeled_block("Animatiestijl-aanpak",
    "Cartoon-flat met dikke outlines, karakter-gedreven, volledig in NESAclinics huisstijl-kleuren. Géén turquoise (referentievideo-kleur). Géén vonkjes of bliksem-symbolen — bewuste keuze om de elektriciteit-angst weg te nemen.")
labeled_block("Bewegingsstijl",
    "Rustig, organisch tempo (ease-in/ease-out). Idle-animaties bij karakters (ademen, lichte beweging) voor levendigheid zonder afleiding.")
labeled_block("Geluid / muziek",
    "Zachte instrumentale achtergrondmuziek. Subtiele sfx bij pulse-animaties (scenes 4, 5, 6). Géén schokkende of elektrische geluidseffecten.")
labeled_block("Lettertype on-screen",
    "Open Sans (of huisstijl-equivalent). Headlines in Bold, body in Regular. Tekst minimaal 1.5 sec in beeld.")
labeled_block("Eindframe-duur",
    "Logo, CTA, URL en QR-code minimaal 2 seconden stilstaand zichtbaar.")
labeled_block("Logo-positionering",
    "NESAclinics logo subtiel rechtsboven vanaf scene 4 tot eind (kleine watermark). In eindframe (scene 10) groot centraal.")
labeled_block("Persona-continuïteit",
    "De drie personages (Peter, Marieke, Sophie) komen terug in scene 1 (klachten) en scene 8 (hersteld). Marieke is de centrale figuur in scene 2 (probleem) en scene 6 (sessie) om herkenbaarheid te versterken.")

# ── Wijzigingsoverzicht ──
H1("Wijzigingsoverzicht t.o.v. V1")
para("Verzameld overzicht van alle aanpassingen na klant-review en meeting op 26-05.")
wijzigingen = [
    ("1", "Drie generieke karakters → drie data-driven persona's (Peter, Marieke, Sophie)"),
    ("2", "Voice-over: 'oorzaak bleef onduidelijk' → 'klachten blijven bestaan'"),
    ("3", "Celniveau/Na-K-Ca verhaal → centraal zenuwstelsel in 'stand-by' (NESA-eigen taal)"),
    ("4", "Apparaat-beeld: handvat met ronde kop → wit kastje met 5 gekleurde aansluitingen, kabels naar handschoenen en sokken"),
    ("4", "Voice-over uitgebreid met 3 voordelen (pijnloos, niet-invasief, ondersteunt zelfherstel)"),
    ("4-5", "Vonkjes/elektriciteit-symbolen weggehaald — vervangen door zachte pulse-flow"),
    ("5", "WERKING-scene: nu met 4 hersengebieden uit NESA-infographic (cortex, limbisch, cerebellum, hersenstam) + zelfregulerend systeem"),
    ("6", "Sessie: 'zit ontspannen' → 'lig ontspannen op behandelbank', 'volledig pijnvrij' expliciet benoemd"),
    ("6", "Therapeut: leeftijd ~40 (niet 30). Patiënt: liggend half-onderuit op behandelbank, sessieduur ca. 1 uur"),
    ("7", "NIEUWE scene Behandelaars: fysiotherapeut, osteopaat, arts — speciaal opgeleid in NESA-therapie. Locaties-info (40+) hierin geïntegreerd"),
    ("8", "Voice-over concreet: 'binnen 3 tot 5 sessies eerste verbetering' + concrete resultaten (slaap, pijn, energie, focus)"),
    ("9", "NIEUWE scene Het traject: aanmelding → bellen binnen 7 dagen → intake (€157) → 10 sessies (€925) + verwijzing naar vergoedingsinfo"),
    ("10", "CTA: 'Vrijblijvend kennismaking' geschrapt (bestaat niet bij NESAclinics) → 'Plan je intake'"),
    ("Algemeen", "Tone-of-voice verschoven van marketing/pijn-druk naar informatief — de video wordt getoond aan al-geïnteresseerde kijkers, niet als cold ad"),
    ("Algemeen", "Duur 60s → 1:10 om alle uitleg over werking, sessie, behandelaars en kosten kwijt te kunnen"),
]
tbl2 = doc.add_table(rows=len(wijzigingen) + 1, cols=2)
tbl2.style = "Light Grid Accent 1"
hdr = tbl2.rows[0].cells
hdr[0].text = "Scene"
hdr[1].text = "Wijziging"
for c in hdr:
    for r in c.paragraphs[0].runs:
        r.bold = True
for i, (scene, wijz) in enumerate(wijzigingen):
    cells = tbl2.rows[i + 1].cells
    cells[0].text = scene
    cells[1].text = wijz

doc.save(OUT)
print(f"Saved: {OUT}")
